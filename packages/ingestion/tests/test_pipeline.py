from pathlib import Path

import fitz
import pytest
from docx import Document
from lexguard_ingestion.config import IngestionSettings
from lexguard_ingestion.pipeline import IngestionPipeline


@pytest.fixture
def pipeline() -> IngestionPipeline:
    return IngestionPipeline(IngestionSettings(ingestion_ocr_min_chars_per_page=30))


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Article 1", fontsize=18)
    page.insert_text((72, 110), "This agreement governs the parties.", fontsize=12)
    page.insert_text((72, 150), "Article 2", fontsize=18)
    page.insert_text((72, 188), "Termination clauses apply herein.", fontsize=12)
    doc.save(path)
    doc.close()
    return path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Chapter 1", level=1)
    document.add_paragraph("Opening content for the chapter.")
    document.add_heading("Section A", level=2)
    document.add_paragraph("Nested section content.")
    document.save(path)
    return path


@pytest.mark.asyncio
async def test_process_pdf(pipeline: IngestionPipeline, sample_pdf: Path) -> None:
    from lexguard_ingestion.models import DocumentFormat

    result = await pipeline.process("doc-001", sample_pdf, DocumentFormat.PDF)

    assert result.document_id == "doc-001"
    assert len(result.sections) >= 1
    assert any("Article" in s.title or "agreement" in s.content for s in result.sections)


@pytest.mark.asyncio
async def test_process_docx(pipeline: IngestionPipeline, sample_docx: Path) -> None:
    from lexguard_ingestion.models import DocumentFormat

    result = await pipeline.process("doc-002", sample_docx, DocumentFormat.DOCX)

    assert result.document_id == "doc-002"
    assert result.sections[0].title == "Chapter 1"
    assert len(result.sections[0].subsections) == 1
    assert result.sections[0].subsections[0].title == "Section A"
